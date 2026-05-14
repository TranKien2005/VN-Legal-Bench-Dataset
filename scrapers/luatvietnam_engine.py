import requests
from bs4 import BeautifulSoup
import docx
import io
import json
import time
import random
import pdfplumber
from pathlib import Path
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor

# Import config
import sys
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from config.settings import settings

class LuatVietnamEngine:
    def __init__(self, num_workers=3):
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Referer": "https://luatvietnam.vn/ban-an/tim-ban-an.html"
        })
        self.base_url = "https://luatvietnam.vn"
        self.num_workers = num_workers
        self.raw_dir = settings.RAW_DIR / "court_cases"
        self.raw_dir.mkdir(parents=True, exist_ok=True)
        self.min_char_count = 1000  # Ngưỡng tối thiểu 1000 chữ theo yêu cầu

    def extract_docx_text(self, content_bytes):
        """Trích xuất text từ file docx trong RAM."""
        try:
            doc = docx.Document(io.BytesIO(content_bytes))
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return "\n".join(full_text)
        except Exception as e:
            print(f"  [LỖI] Không thể đọc docx: {e}")
            return ""

    def extract_pdf_text(self, content_bytes):
        """Trích xuất text từ file PDF trong RAM bằng pdfplumber."""
        try:
            full_text = []
            with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text.append(text)
            return "\n".join(full_text)
        except Exception as e:
            print(f"  [LỖI] Không thể đọc PDF: {e}")
            return ""

    def scrape_detail_page(self, detail_url):
        """Cào trang chi tiết (Thuộc tính, Tóm tắt và file nội dung)."""
        try:
            # Tạo header cục bộ cho thread này
            headers = self.session.headers.copy()
            headers["Referer"] = detail_url
            
            response = self.session.get(detail_url, headers=headers, timeout=30)
            if response.status_code != 200:
                print(f"  [LỖI] Không thể truy cập trang chi tiết {detail_url} (Status {response.status_code})")
                return None

            soup = BeautifulSoup(response.text, 'html.parser')
            doc_record = {
                "url": detail_url,
                "scraped_at": datetime.now().isoformat(),
                "metadata": {
                    "summary": "",
                    "pdf_url": None,
                    "docx_url": None
                },
                "raw_text": ""
            }

            # 1. Lấy Metadata từ bảng Thuộc tính
            attr_table = soup.find('table', class_='table-thuoc-tinh')
            if not attr_table:
                attr_table = soup.find('table', class_='table-attribute')
            if not attr_table:
                attr_table = soup.find('table')

            # Initialize all required metadata fields with None
            doc_record["metadata"].update({
                "Số hiệu": None,
                "Ngày ban hành": None,
                "Lĩnh vực": None,
                "Tòa án xét xử": None,
                "Quan hệ pháp luật": None,
                "Cấp xét xử": None,
                "Loại văn bản": None,
                "Người ký": None,
                "summary": None,
                "pdf_url": None,
                "docx_url": None,
                "source_format": None
            })

            if attr_table:
                for row in attr_table.find_all('tr'):
                    cells = row.find_all(['td', 'th'])
                    if len(cells) >= 2:
                        key = cells[0].get_text(strip=True).replace(':', '')
                        val = cells[1].get_text(separator=' ', strip=True)
                        doc_record["metadata"][key] = val

            # Hậu kiểm: Đảm bảo chỉ lấy bản án Sơ thẩm
            judicial_level = doc_record["metadata"].get("Cấp xét xử", "")
            if judicial_level and "Sơ thẩm" not in judicial_level:
                print(f"  [!] Bỏ qua bản án: Cấp xét xử là '{judicial_level}', không phải Sơ thẩm.")
                return None
            
            # 2. Lấy Tóm tắt (Tóm tắt Bản án)
            import re
            summary_label = soup.find(lambda tag: tag.name == "div" and "doc-headding" in tag.get("class", []) and re.search(r"Tóm tắt Bản án", tag.text, re.I))
            if summary_label:
                # Tìm block tiếp theo có chứa document-body
                summary_block = summary_label.find_parent('div', class_='block').find_next_sibling('div', class_='block')
                if summary_block:
                    body = summary_block.find('div', class_='document-body')
                    if body:
                        doc_record["metadata"]["summary"] = body.get_text(strip=True)

            # 3. Tìm link tải file (DOCX và PDF)
            download_blocks = soup.find_all('div', class_='list-download')
            for block in download_blocks:
                a = block.find('a', href=True)
                if a:
                    href = a['href']
                    if not href.startswith('http'):
                        href = self.base_url + href
                    
                    if href.lower().endswith('.docx'):
                        doc_record["metadata"]["docx_url"] = href
                    elif href.lower().endswith('.pdf'):
                        doc_record["metadata"]["pdf_url"] = href

            # 4. Xử lý lấy nội dung văn bản (raw_text) với fallback
            final_text = ""
            source_used = None

            # Thử DOCX trước
            if doc_record["metadata"]["docx_url"]:
                print(f"  -> Thử tải DOCX: {doc_record['metadata']['docx_url']}")
                try:
                    headers = self.session.headers.copy()
                    headers["Referer"] = detail_url
                    res = self.session.get(doc_record["metadata"]["docx_url"], headers=headers, timeout=60)
                    if res.status_code == 200:
                        text = self.extract_docx_text(res.content)
                        if len(text) >= self.min_char_count:
                            final_text = text
                            source_used = "docx"
                    else:
                        print(f"  [!] Lỗi tải DOCX (Status {res.status_code})")
                except Exception as e:
                    print(f"  [!] Lỗi tải/đọc DOCX: {e}")

            # Fallback sang PDF nếu DOCX thất bại hoặc không đủ dài
            if len(final_text) < self.min_char_count and doc_record["metadata"]["pdf_url"]:
                print(f"  -> Thử tải PDF (fallback): {doc_record['metadata']['pdf_url']}")
                try:
                    headers = self.session.headers.copy()
                    headers["Referer"] = detail_url
                    res = self.session.get(doc_record["metadata"]["pdf_url"], headers=headers, timeout=60)
                    if res.status_code == 200:
                        text = self.extract_pdf_text(res.content)
                        if len(text) >= self.min_char_count:
                            final_text = text
                            source_used = "pdf"
                    else:
                        print(f"  [!] Lỗi tải PDF (Status {res.status_code})")
                except Exception as e:
                    print(f"  [!] Lỗi tải/đọc PDF: {e}")

            if len(final_text) < self.min_char_count:
                print(f"  [!] Bỏ qua bản án: Nội dung quá ngắn ({len(final_text)} ký tự).")
                return None  # Bỏ qua hoàn toàn vụ án này

            doc_record["raw_text"] = final_text
            doc_record["metadata"]["source_format"] = source_used
            
            print(f"  -> Thành công: Lấy được {len(final_text)} ký tự từ {source_used}.")
            return doc_record
        except Exception as e:
            print(f"  [LỖI] Lỗi trang chi tiết {detail_url}: {e}")
            return None

    def scrape_search_page(self, params):
        """Lấy danh sách link và tiêu đề bản án từ trang tìm kiếm."""
        try:
            search_url = f"{self.base_url}/ban-an/tim-ban-an.html"
            response = self.session.get(search_url, params=params, timeout=30)
            if response.status_code != 200:
                print(f"  [LỖI] Không thể access trang tìm kiếm (Status {response.status_code})")
                return []

            soup = BeautifulSoup(response.text, 'html.parser')
            results = []
            title_containers = soup.find_all('h3', class_='entry-title')
            for h3 in title_containers:
                a = h3.find('a', href=True)
                if a:
                    full_link = a['href']
                    if not full_link.startswith('http'):
                        full_link = self.base_url + full_link
                    
                    # Chỉ lấy các link bản án
                    if "/ban-an/" in full_link:
                        results.append({
                            "url": full_link,
                            "title": a.get_text(separator=' ', strip=True)
                        })
            
            if not results:
                # Thử selector khác rộng hơn
                items = soup.select('div.doc-item h3 a[href]')
                for a in items:
                    href = a['href']
                    if "/ban-an/" in href:
                         results.append({
                            "url": self.base_url + href if not href.startswith('http') else href,
                            "title": a.get_text(separator=' ', strip=True)
                        })

            print(f"  -> Tìm thấy {len(results)} bản án trên trang này.")
            return results

        except Exception as e:
            print(f"  [LỖI] Lỗi trang tìm kiếm: {e}")
            return []

    def run(self, max_pages=1, start_page=1, custom_params=None):
        """Chạy quy trình cào dữ liệu."""
        default_params = {
            "SearchKeyword": "",
            "DateFrom": "01/01/2010",
            "DateTo": datetime.now().strftime("%d/%m/%Y"),
            "IsSearchExact": "0",
            "SearchByDate": "publishDate",
            "CourtLevelId": "0",
            "CourtId": "0",
            "JudicialLevelId": "1",
            "LawJudgTypeId": "1",
            "CateId": "0",
            "LegalRelationId": "0",
            "TypeSearch": "0"
        }
        
        if custom_params:
            default_params.update(custom_params)

        print(f"[*] Tham số tìm kiếm: {default_params}")

        # Mồi cookie trang chủ trước
        self.session.get(self.base_url)

        for page in range(start_page, start_page + max_pages):
            print(f"\n[*] Đang cào trang {page}...")
            page_params = default_params.copy()
            page_params["Page"] = page
            
            links = self.scrape_search_page(page_params)
            if not links:
                print(f"  [!] Không tìm thấy thêm liên kết nào ở trang {page}. Dừng.")
                break
            
            # Xử lý các link song song bằng ThreadPoolExecutor
            page_results = []
            with ThreadPoolExecutor(max_workers=self.num_workers) as executor:
                # Map detail page scraping tasks
                future_to_url = {executor.submit(self.scrape_detail_page, link["url"]): link["url"] for link in links}
                
                for future in future_to_url:
                    result = future.result()
                    if result:
                        page_results.append(result)
                    # Giảm tải: nghỉ nhẹ giữa các kết quả nhận được (tùy chọn)
                    # time.sleep(random.uniform(0.1, 0.5))
            
            if page_results:
                # Tạo tên file gợi nhớ dựa trên tham số
                kw = default_params.get("SearchKeyword") or "none"
                kw_clean = "".join([c if c.isalnum() else "_" for c in kw])
                d_from = default_params.get("DateFrom", "").replace("/", "")
                d_to = default_params.get("DateTo", "").replace("/", "")
                
                filename = f"luatvietnam_K-{kw_clean}_F-{d_from}_T-{d_to}_P-{page}.json"
                output_path = self.raw_dir / filename
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    json.dump(page_results, f, ensure_ascii=False, indent=2)
                print(f"--- Đã lưu {len(page_results)} bản án vào {filename} ---")
            
            if page < start_page + max_pages - 1:
                wait_time = random.uniform(3, 7)
                print(f"[*] Nghỉ {wait_time:.1f} giây trước khi sang trang tiếp theo...")
                time.sleep(wait_time)

        print("\n[KẾT THÚC] Hoàn tất quá trình cào dữ liệu.")


class LuatVietnamLegalDocEngine:
    def __init__(self, num_workers=3, min_char_count=200):
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Referer": "https://luatvietnam.vn/van-ban/tim-van-ban.html"
        })
        self.base_url = "https://luatvietnam.vn"
        self.num_workers = num_workers
        self.min_char_count = min_char_count
        self.raw_dir = settings.RAW_DIR / "legal_docs"
        self.raw_dir.mkdir(parents=True, exist_ok=True)

    def _clean_text(self, text):
        if not text:
            return None
        cleaned = " ".join(text.split())
        return cleaned if cleaned and cleaned.lower() not in {"đã biết", "đang cập nhật"} else None

    def _full_url(self, href):
        if not href:
            return None
        if href.startswith("http"):
            return href
        if not href.startswith("/"):
            href = "/" + href
        return self.base_url + href

    def extract_docx_text(self, content_bytes):
        try:
            doc = docx.Document(io.BytesIO(content_bytes))
            return "\n".join(para.text for para in doc.paragraphs)
        except Exception as e:
            print(f"  [LỖI] Không thể đọc DOC/DOCX: {e}")
            return ""

    def extract_pdf_text(self, content_bytes):
        try:
            full_text = []
            with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text.append(text)
            return "\n".join(full_text)
        except Exception as e:
            print(f"  [LỖI] Không thể đọc PDF: {e}")
            return ""

    def _extract_overview_table(self, soup):
        container = soup.select_one("#tomtat") or soup
        table = container.find("table")
        fields = {}
        if not table:
            return fields

        for row in table.find_all("tr"):
            cells = row.find_all(["td", "th"], recursive=False)
            if len(cells) < 2:
                continue
            for idx in range(0, len(cells) - 1, 2):
                label = self._clean_text(cells[idx].get_text(" ", strip=True).replace(":", ""))
                if not label:
                    continue
                if label in {"Tình trạng hiệu lực", "Áp dụng"}:
                    continue
                value_cell = cells[idx + 1]
                if label == "Lĩnh vực":
                    values = [a.get_text(" ", strip=True) for a in value_cell.find_all("a")]
                    value = ", ".join(v for v in values if v) or value_cell.get_text(" ", strip=True)
                else:
                    value = value_cell.get_text(" ", strip=True)
                cleaned_value = self._clean_text(value)
                if cleaned_value:
                    fields[label] = cleaned_value
        return fields

    def _extract_summary(self, soup):
        container = soup.select_one("#tomtat") or soup
        summary = container.select_one(".doc-summary")
        if summary:
            return self._clean_text(summary.get_text("\n", strip=True))
        heading = container.find(lambda tag: tag.name in {"h2", "h3", "div"} and "tóm tắt" in tag.get_text(" ", strip=True).lower())
        if heading:
            body = heading.find_next(lambda tag: tag.name == "div" and ("document-body" in tag.get("class", []) or "doc-summary" in tag.get("class", [])))
            if body:
                return self._clean_text(body.get_text("\n", strip=True))
        meta = soup.find("meta", attrs={"name": "description"})
        return self._clean_text(meta.get("content")) if meta else None

    def _extract_download_links(self, soup):
        container = soup.select_one("#tomtat") or soup
        links = []
        for a in container.select(".list-download a[href], a[href]"):
            href = self._full_url(a.get("href"))
            if not href:
                continue
            label = self._clean_text(a.get_text(" ", strip=True)) or ""
            title = self._clean_text(a.get("title")) or ""
            marker = f"{label} {title} {href}".lower()
            if not any(x in marker for x in ["pdf", "doc", "word", "tai-file", "tải"]):
                continue
            if "pdf" in marker:
                file_type = "pdf"
            elif any(x in marker for x in ["docx", "doc", "word"]):
                file_type = "doc"
            else:
                file_type = "unknown"
            links.append({
                "file_type": file_type,
                "label": label,
                "title": title,
                "url": href
            })
        deduped = []
        seen = set()
        for link in links:
            if link["url"] in seen:
                continue
            seen.add(link["url"])
            deduped.append(link)
        return deduped

    def _download_raw_text(self, detail_url, download_links):
        headers = self.session.headers.copy()
        headers["Referer"] = detail_url
        ordered = sorted(download_links, key=lambda link: 0 if link.get("file_type") == "doc" else 1)
        for link in ordered:
            file_type = link.get("file_type")
            url = link.get("url")
            if file_type not in {"doc", "pdf"} or not url:
                continue
            print(f"  -> Thử tải {file_type.upper()}: {url}")
            try:
                res = self.session.get(url, headers=headers, timeout=60)
                if res.status_code != 200:
                    print(f"  [!] Lỗi tải {file_type.upper()} (Status {res.status_code})")
                    continue
                text = self.extract_docx_text(res.content) if file_type == "doc" else self.extract_pdf_text(res.content)
                if file_type == "doc" and text.strip():
                    return text, file_type, url
                if file_type == "pdf" and len(text) >= self.min_char_count:
                    return text, file_type, url
            except Exception as e:
                print(f"  [!] Lỗi tải/đọc {file_type.upper()}: {e}")
        return "", None, None

    def scrape_detail_page(self, detail_url, title_web=None):
        try:
            headers = self.session.headers.copy()
            headers["Referer"] = detail_url
            response = self.session.get(detail_url, headers=headers, timeout=30)
            if response.status_code != 200:
                print(f"  [LỖI] Không thể truy cập trang chi tiết {detail_url} (Status {response.status_code})")
                return None

            soup = BeautifulSoup(response.text, "html.parser")
            canonical = soup.find("link", rel="canonical")
            source_url = canonical.get("href") if canonical and canonical.get("href") else detail_url
            title_tag = soup.select_one("#tomtat h1.the-document-title") or soup.find("h1")
            final_title_web = self._clean_text(title_tag.get_text(" ", strip=True)) if title_tag else title_web

            record = {
                "source": "luatvietnam_vanban",
                "url": source_url,
                "title_web": final_title_web,
                "scraped_at": datetime.now().isoformat(),
                "raw_text": ""
            }
            record.update(self._extract_overview_table(soup))
            summary = self._extract_summary(soup)
            if summary:
                record["summary"] = summary
                record["Tóm tắt"] = summary

            download_links = self._extract_download_links(soup)
            record["download_links"] = download_links
            raw_text, source_format, source_doc_url = self._download_raw_text(source_url, download_links)
            if len(raw_text) < self.min_char_count:
                body = soup.select_one("#noidung .the-document-body[data-role='content-body']") or soup.select_one("#noidung .the-document-body")
                if body:
                    html_text = body.get_text("\n", strip=True)
                    if len(html_text) >= self.min_char_count:
                        raw_text = html_text
                        source_format = "html"
                        source_doc_url = source_url

            if len(raw_text) < self.min_char_count:
                print(f"  [!] Bỏ qua văn bản: Nội dung quá ngắn ({len(raw_text)} ký tự).")
                return None

            record["raw_text"] = raw_text
            record["source_format"] = source_format
            record["source_doc_url"] = source_doc_url
            print(f"  -> Thành công: Lấy được {len(raw_text)} ký tự từ {source_format}.")
            return record
        except Exception as e:
            print(f"  [LỖI] Lỗi trang chi tiết {detail_url}: {e}")
            return None

    def scrape_search_page(self, params):
        try:
            response = self.session.get(f"{self.base_url}/van-ban/tim-van-ban.html", params=params, timeout=30)
            if response.status_code != 200:
                print(f"  [LỖI] Không thể access trang tìm kiếm (Status {response.status_code})")
                return []
            soup = BeautifulSoup(response.text, "html.parser")
            results = []
            seen = set()
            for article in soup.select("article.art-search"):
                title_anchor = article.select_one("h2.doc-title a[href]") or article.select_one("h3.entry-title a[href]")
                if not title_anchor:
                    continue
                href = self._full_url(title_anchor.get("href"))
                if not href or href in seen or "/ban-an/" in href:
                    continue
                if not ("-d1.html" in href or "-d10.html" in href):
                    continue
                seen.add(href)
                results.append({"url": href, "title": self._clean_text(title_anchor.get_text(" ", strip=True))})
            print(f"  -> Tìm thấy {len(results)} văn bản trên trang này.")
            return results
        except Exception as e:
            print(f"  [LỖI] Lỗi trang tìm kiếm văn bản: {e}")
            return []

    def run(self, doc_type_ids, output_prefix, max_pages=1, start_page=1, custom_params=None):
        default_params = {
            "PagSize": "100",
            "PageSize": "100",
        }
        if custom_params:
            default_params.update(custom_params)

        self.session.get(self.base_url, timeout=30)
        page_count = 9999 if max_pages == "auto" else int(max_pages)
        for page in range(start_page, start_page + page_count):
            print(f"\n[*] Đang cào trang văn bản {page}...")
            params = default_params.copy()
            params["PageIndex"] = page
            query_params = list(params.items())
            for doc_type_id in doc_type_ids:
                query_params.append(("DocTypeIds", str(doc_type_id)))
            links = self.scrape_search_page(query_params)
            if not links:
                print(f"  [!] Không tìm thấy thêm liên kết nào ở trang {page}. Dừng.")
                break

            page_results = []
            with ThreadPoolExecutor(max_workers=self.num_workers) as executor:
                futures = [executor.submit(self.scrape_detail_page, link["url"], link.get("title")) for link in links]
                for future in futures:
                    result = future.result()
                    if result:
                        page_results.append(result)

            if page_results:
                filename = f"luatvietnam_{output_prefix}_page_{page}.json"
                with open(self.raw_dir / filename, "w", encoding="utf-8") as f:
                    json.dump(page_results, f, ensure_ascii=False, indent=2)
                print(f"--- Đã lưu {len(page_results)} văn bản vào {filename} ---")

            if page < start_page + page_count - 1:
                time.sleep(random.uniform(2, 5))

    def scrape_special_urls(self, urls, output_name="luatvietnam_special_results.json"):
        results = []
        self.session.get(self.base_url, timeout=30)
        for i, url in enumerate(urls, start=1):
            print(f"[{i}/{len(urls)}] Đang xử lý: {url}")
            record = self.scrape_detail_page(url)
            if record:
                results.append(record)
        if results:
            with open(self.raw_dir / output_name, "w", encoding="utf-8") as f:
                json.dump(results, f, ensure_ascii=False, indent=2)
            print(f"\n[DONE] Đã lưu {len(results)} văn bản vào: {self.raw_dir / output_name}")
